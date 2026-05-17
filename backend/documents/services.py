import os
import re
import copy
from io import BytesIO
from docx import Document
from docx.shared import Pt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from django.core.files.base import ContentFile


class DocumentService:

    @staticmethod
    def fill_template(template_path: str, data: dict) -> BytesIO:
        """
        Shablonni ma'lumotlar bilan to'ldiradi
        {{ism}} -> "Ali Valiyev"
        """
        doc = Document(template_path)

        def replace_in_paragraph(paragraph, data):
            for key, value in data.items():
                placeholder = f'{{{{{key}}}}}'
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))

        # Paragraphlarda almashtirish
        for para in doc.paragraphs:
            replace_in_paragraph(para, data)

        # Jadvallarda almashtirish
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, data)

        # Header/Footer
        for section in doc.sections:
            for para in section.header.paragraphs:
                replace_in_paragraph(para, data)
            for para in section.footer.paragraphs:
                replace_in_paragraph(para, data)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    @staticmethod
    def docx_to_pdf(docx_buffer: BytesIO, data: dict) -> BytesIO:
        """
        DOCX matnini PDF ga o'tkazadi
        (Production'da LibreOffice ishlatish tavsiya etiladi)
        """
        doc = Document(docx_buffer)

        pdf_buffer = BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=A4)
        width, height = A4

        # O'zbek/Kirill uchun font
        try:
            pdfmetrics.registerFont(TTFont('DejaVu', 'DejaVuSans.ttf'))
            font_name = 'DejaVu'
        except:
            font_name = 'Helvetica'

        y = height - 50
        c.setFont(font_name, 11)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                y -= 10
                continue

            # Uzun textni qatorlarga bo'lish
            words = text.split()
            line = ""
            for word in words:
                test_line = f"{line} {word}".strip()
                if c.stringWidth(test_line, font_name, 11) < width - 100:
                    line = test_line
                else:
                    c.drawString(50, y, line)
                    y -= 20
                    line = word

            if line:
                c.drawString(50, y, line)
                y -= 20

            if y < 50:
                c.showPage()
                c.setFont(font_name, 11)
                y = height - 50

        c.save()
        pdf_buffer.seek(0)
        return pdf_buffer
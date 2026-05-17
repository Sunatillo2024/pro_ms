# apps/templates_app/models.py
import re
from django.db import models
from django.conf import settings


class DocumentTemplate(models.Model):
    name = models.CharField(max_length=255)
    description = models.TextField(blank=True)
    file = models.FileField(upload_to='templates/')
    created_by = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.SET_NULL,
        null=True
    )
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    is_active = models.BooleanField(default=True)

    def get_placeholders(self):
        """Shablondagi {{placeholder}} larni topadi"""
        from docx import Document
        doc = Document(self.file.path)
        placeholders = set()

        for para in doc.paragraphs:
            found = re.findall(r'\{\{(\w+)\}\}', para.text)
            placeholders.update(found)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    found = re.findall(r'\{\{(\w+)\}\}', cell.text)
                    placeholders.update(found)

        return list(placeholders)

    def __str__(self):
        return self.name

    class Meta:
        verbose_name = "Shablon"
        verbose_name_plural = "Shablonlar"